import os
import time
from functools import wraps
from flask import Flask, request, jsonify, send_file
import mysql.connector
from flask_cors import CORS
from dotenv import load_dotenv
from openai import OpenAI
import requests   # ✅ needed for Flutterwave

# 📂 File generation imports (Day 4)
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer  # PDF
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document  # DOCX
import pandas as pd        # CSV + Excel
from openpyxl import Workbook  # XLSX

# --- Load environment variables ---
load_dotenv()
print("AI API KEY:", os.getenv("AI_API_KEY")[:10])  # shows first 10 chars only

app = Flask(__name__)
CORS(app, resources={r"/api/*": {"origins": ["http://127.0.0.1:8000", "http://localhost:8000"]}})

# --- OpenAI setup ---
client = OpenAI(api_key=os.getenv("AI_API_KEY"))

# --- Cache & Rate limiting ---
CACHE = {}
CACHE_TTL = 60 * 60  # 1 hour
RATE_LIMIT = {}
RATE_WINDOW = 60  # seconds
RATE_MAX = 30

def rate_limit(func):
    @wraps(func)
    def wrapper(*args, **kwargs):
        ip = request.remote_addr
        now = int(time.time())
        window = now // RATE_WINDOW
        key = f"{ip}:{window}"
        RATE_LIMIT.setdefault(key, 0)
        if RATE_LIMIT[key] >= RATE_MAX:
            return jsonify({"status": "error", "message": "Rate limit exceeded"}), 429
        RATE_LIMIT[key] += 1
        return func(*args, **kwargs)
    return wrapper

# --- Database connection ---
DB_CONFIG = {
    "user": os.environ.get("DB_USER"),
    "password": os.environ.get("DB_PASSWORD"),
    "host": os.environ.get("DB_HOST"),
    "database": os.environ.get("DB_NAME"),
}

def db_conn():
    return mysql.connector.connect(**DB_CONFIG)

# --- Routes ---
@app.route("/")
def home():
    return "Flask is connected to MySQL, OpenAI, and Flutterwave!"

@app.route('/api/students')
@rate_limit
def students():
    conn = db_conn()
    cur = conn.cursor(dictionary=True)
    cur.execute('SELECT id, reg_no, first_name, last_name, class FROM students LIMIT 200')
    rows = cur.fetchall()
    cur.close(); conn.close()
    return jsonify(rows)

@app.route('/api/assessments', methods=['GET', 'POST'])
@rate_limit
def assessments():
    conn = db_conn()
    cur = conn.cursor(dictionary=True)

    if request.method == 'POST':
        data = request.json
        cur.execute(
            "INSERT INTO assessments (title, description, competency, max_score) VALUES (%s,%s,%s,%s)",
            (data.get('title'), data.get('description'), data.get('competency'), data.get('max_score', 100))
        )
        conn.commit()
        new_id = cur.lastrowid
        cur.close(); conn.close()
        return jsonify({'status': 'success', 'id': new_id}), 201

    # GET request
    cur.execute('SELECT * FROM assessments')
    rows = cur.fetchall()
    cur.close(); conn.close()
    return jsonify(rows)

# --- AI Lesson Plan ---
@app.route('/api/ai/lesson_plan', methods=['POST', 'OPTIONS'])
def lesson_plan():
    if request.method == 'OPTIONS':
        return jsonify({"status": "ok"}), 200

    try:
        data = request.get_json(force=True) or {}
        subject = data.get("subject", "Unknown Subject")
        grade = data.get("grade", "N/A")
        competency = data.get("competency", "General Competency")
        duration = data.get("duration", "40 minutes")

        cache_key = f"lessonplan-{subject}-{grade}-{competency}-{duration}"
        if cache_key in CACHE and (time.time() - CACHE[cache_key]["time"]) < CACHE_TTL:
            ai_plan = CACHE[cache_key]["data"]
        else:
            prompt = (
                f"Generate a detailed lesson plan for {subject}, {grade}, aligned with the Kenya CBC curriculum."
                f" Competency: {competency}. Duration: {duration}.\n\n"
                "Return the plan in 4 sections:\n1. Objectives\n2. Materials\n3. Activities\n4. Assessment\n"
                "Keep each section short and practical."
            )
            response = client.chat.completions.create(
                model=os.getenv("LLM_MODEL", "gpt-4o-mini"),
                messages=[
                    {"role": "system", "content": "You are a helpful teacher assistant for Kenya’s CBC curriculum."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.6,
                max_tokens=400
            )
            ai_plan = response.choices[0].message.content.strip()
            CACHE[cache_key] = {"data": ai_plan, "time": time.time()}

        return jsonify({
            "status": "success",
            "message": "Lesson plan generated successfully!",
            "subject": subject,
            "grade": grade,
            "competency": competency,
            "duration": duration,
            "plan": ai_plan
        }), 200

    except Exception as e:
        error_msg = str(e)
        if "insufficient_quota" in error_msg or "You exceeded your current quota" in error_msg:
            return jsonify({"status": "error", "message": "⚠️ OpenAI quota exceeded. Check billing."}), 429
        return jsonify({"status": "error", "message": error_msg}), 500

# --- AI Parent Note ---
@app.route('/api/ai/parent_note', methods=['POST', 'OPTIONS'])
def parent_note():
    if request.method == 'OPTIONS':
        return jsonify({"status": "ok"}), 200

    try:
        data = request.get_json(force=True) or {}
        student_name = data.get("student_name", "Student")
        progress = data.get("progress", "")
        lang = data.get("lang", "en")

        cache_key = f"parentnote-{student_name}-{progress}-{lang}"
        if cache_key in CACHE and (time.time() - CACHE[cache_key]["time"]) < CACHE_TTL:
            note_text = CACHE[cache_key]["data"]
        else:
            prompt = f"Write a short {lang.upper()} parent note for {student_name}. Progress report: {progress}."
            response = client.chat.completions.create(
                model=os.getenv("LLM_MODEL", "gpt-4o-mini"),
                messages=[{"role": "user", "content": prompt}],
                temperature=0.5,
                max_tokens=150
            )
            note_text = response.choices[0].message.content
            CACHE[cache_key] = {"data": note_text, "time": time.time()}

        return jsonify({"status": "success", "student_name": student_name, "note": note_text}), 200

    except Exception as e:
        error_msg = str(e)
        if "insufficient_quota" in error_msg or "You exceeded your current quota" in error_msg:
            return jsonify({"status": "error", "message": "⚠️ OpenAI quota exceeded. Check billing."}), 429
        return jsonify({"status": "error", "message": error_msg}), 500

# --- Export / Download Routes ---
@app.route('/api/export/pdf')
def export_pdf():
    try:
        filename = "student_report.pdf"
        doc = SimpleDocTemplate(filename)
        styles = getSampleStyleSheet()
        story = [Paragraph("Student Report", styles['Title']), Spacer(1, 12)]
        story.append(Paragraph("This is a generated PDF report for students.", styles['Normal']))
        doc.build(story)
        return send_file(filename, as_attachment=True)
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/api/export/docx')
def export_docx():
    try:
        filename = "student_report.docx"
        doc = Document()
        doc.add_heading("Student Report", level=1)
        doc.add_paragraph("This is a generated Word (DOCX) report for students.")
        doc.save(filename)
        return send_file(filename, as_attachment=True)
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/api/export/csv')
def export_csv():
    try:
        filename = "student_report.csv"
        data = {"Name": ["Alice", "Bob"], "Grade": ["A", "B"]}
        df = pd.DataFrame(data)
        df.to_csv(filename, index=False)
        return send_file(filename, as_attachment=True)
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/api/export/excel')
def export_excel():
    try:
        filename = "student_report.xlsx"
        wb = Workbook()
        ws = wb.active
        ws.title = "Students"
        ws.append(["Name", "Grade"])
        ws.append(["Alice", "A"])
        ws.append(["Bob", "B"])
        wb.save(filename)
        return send_file(filename, as_attachment=True)
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

# --- Flutterwave Payment Route ---
@app.route('/api/pay', methods=['POST'])
def initiate_payment():
    try:
        data = request.json
        amount = data.get("amount")
        email = data.get("email")
        name = data.get("name")

        url = "https://api.flutterwave.com/v3/payments"
        headers = {
            "Authorization": f"Bearer {os.getenv('FLW_SECRET_KEY')}",
            "Content-Type": "application/json"
        }
        payload = {
            "tx_ref": f"SAINTCBC-{name}-{int(time.time())}",
            "amount": amount,
            "currency": "KES",
            "redirect_url": "http://localhost:3000/payment-status",
            "customer": {
                "email": email,
                "name": name
            },
            "customizations": {
                "title": "Saint CBC School Fees",
                "description": "Pay school fees securely via Flutterwave",
                "logo": "https://your-logo-link.png"
            }
        }

        r = requests.post(url, json=payload, headers=headers)
        return jsonify(r.json())

    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


# --- Run the app ---
if __name__ == "__main__":
    app.run(debug=True)
